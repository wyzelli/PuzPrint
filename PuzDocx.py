from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # For qualified names

# Configuration
CELL_SIZE_CM = 0.8

def draw_sudoku_table(doc, puzzle):
    """Create a Sudoku table with proper borders and centered numbers"""
    table = doc.add_table(rows=9, cols=9)
    table.style = 'Table Grid'

    for row_idx in range(9):
        row = table.rows[row_idx]
        for col_idx in range(9):
            cell = row.cells[col_idx]
            cell.width = Cm(CELL_SIZE_CM)
            
            # Add number if not zero
            digit = puzzle[row_idx * 9 + col_idx]
            if digit != '0':
                p = cell.paragraphs[0]
                run = p.add_run(digit)
                run.bold = True
                run.font.size = Pt(14)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Set custom borders
            set_cell_borders(cell, row_idx, col_idx)

def set_cell_borders(cell, row_idx, col_idx):
    """Set custom borders for 3x3 Sudoku grid"""
    tc_pr = cell._tc.get_or_add_tcPr()
    
    # Create borders element with proper namespace
    borders = OxmlElement(qn('w:tcBorders'))

    # Border definitions using qualified names
    border_types = [
        (qn('w:top'), 4),
        (qn('w:left'), 4),
        (qn('w:bottom'), 4),
        (qn('w:right'), 4)
    ]

    # Create thin borders first
    for border_name, size in border_types:
        border = OxmlElement(qn('w:top')) if border_name == qn('w:top') else OxmlElement(border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        borders.append(border)

    # Thicken borders between 3x3 boxes
    if row_idx % 3 == 0:  # Top of 3x3 box
        borders.xpath(qn('w:top'))[0].set(qn('w:sz'), '12')
    if (row_idx + 1) % 3 == 0:  # Bottom of 3x3 box
        borders.xpath(qn('w:bottom'))[0].set(qn('w:sz'), '12')
    if col_idx % 3 == 0:  # Left of 3x3 box
        borders.xpath(qn('w:left'))[0].set(qn('w:sz'), '12')
    if (col_idx + 1) % 3 == 0:  # Right of 3x3 box
        borders.xpath(qn('w:right'))[0].set(qn('w:sz'), '12')

    # Replace existing borders or add new ones
    existing_borders = tc_pr.xpath('.//w:tcBorders')
    if existing_borders:
        tc_pr.remove(existing_borders[0])
    tc_pr.append(borders)

# Rest of the code remains the same
def load_puzzles(file_path):
    puzzles = []
    with open(file_path, 'r') as f:
        for line in f:
            clean_line = line.strip().replace(" ", "")[:81]
            if len(clean_line) == 81 and clean_line.isdigit():
                puzzles.append(clean_line)
    return puzzles

def create_docx(puzzles, output_file):
    doc = Document()
    for i in range(0, len(puzzles), 2):
        if i > 0:
            doc.add_page_break()
        draw_sudoku_table(doc, puzzles[i])
        if i + 1 < len(puzzles):
            doc.add_paragraph()
            draw_sudoku_table(doc, puzzles[i + 1])
    doc.save(output_file)
    print(f"Generated DOCX with {len(puzzles)} puzzles: {output_file}")

if __name__ == "__main__":
    puzzles = load_puzzles("puzzles.txt")
    if puzzles:
        create_docx(puzzles, "sudoku_puzzles.docx")
    else:
        print("No valid puzzles found in puzzles.txt")
